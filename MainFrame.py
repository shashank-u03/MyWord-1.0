#############################################################################################################################
# Author:- Shashank Upadhyay
# This file is the main GUI file that is to be launched via command - python MainFrame.py
# This file will run on 3.10.5 and updated versions
#############################################################################################################################
from tkinter import *
from PIL import Image, ImageTk
from tkinter import Tk, Frame, Menu
from docx import Document
import docx
from docx.shared import Pt
from tkinter import filedialog
import huffman
import charset_normalizer.md__mypyc
from tkinter import font
from tkinter.filedialog import asksaveasfile
from tkinter.filedialog import askopenfilename
from fpdf import FPDF
import pdfplumber
import PyPDF2
import io
import os

root = Tk()

root.title("MSTC Word")
# Dimensions are width x height
root.geometry("1000x800")
root.config(bg="white")

# Adding a text widget
inputTxt = Text(root, width=100, height=80, bg="white", wrap=WORD)
inputTxt.configure(font=("helvetica", 15))
# inputTxt.insert(END,"My name is Shashank")

# Rightframe for widgets
rightframe = Frame(root, bg="dark blue")
rightframe.pack(side=RIGHT)

# Leftframe for widgets
leftFrame = Frame(root, bg="dark blue")
leftFrame.pack(side=LEFT)

# function to find a word in the text widget


def find():
    # remove tag 'found' from index 1 to END
    inputTxt.tag_remove('found', '1.0', END)

    # returns to widget currently in focus
    s = edit.get()

    if (s):
        idx = '1.0'
        while 1:
            # searches for desried string from index 1
            idx = inputTxt.search(s, idx, nocase=1,
                                  stopindex=END)

            if not idx:
                break
            # last index sum of current index and
            # length of text
            lastidx = '% s+% dc' % (idx, len(s))

            # overwrite 'Found' at idx
            inputTxt.tag_add('found', idx, lastidx)
            idx = lastidx

        # mark located string as red

        inputTxt.tag_config('found', foreground='red')
    edit.focus_set()

# function to find and replace string in text widget

def replace_string():
    # remove tag 'found' from index 1 to END
    inputTxt.tag_remove('found', '1.0', END)

    # returns to widget currently in focus
    s = edit.get()
    r = edit2.get()

    if (s and r):
        idx = '1.0'
        while 1:
            # searches for desried string from index 1
            idx = inputTxt.search(s, idx, nocase=1,
                                  stopindex=END)
            # print(idx)
            if not idx:
                break

            # last index sum of current index and
            # length of text
            lastidx = '% s+% dc' % (idx, len(s))

            inputTxt.delete(idx, lastidx)
            inputTxt.insert(idx, r)

            lastidx = '% s+% dc' % (idx, len(r))

            # overwrite 'Found' at idx
            inputTxt.tag_add('found', idx, lastidx)
            idx = lastidx

        # mark located string as red
        inputTxt.tag_config('found', foreground='blue', background='white')
    edit.focus_set()


# funtion to change font style
def font_style():
    style = clicked.get()
    # print(style)
    font_style_to_be_changed = style
    inputTxt.configure(font=(font_style_to_be_changed, 15))

# function to add image into the text widget

def insert_image():
    global my_image
    imagefile = filedialog.askopenfilename(
        initialdir="/", title="Select a File", filetypes=(("PNG file", "*.png*"), ("JPG file", "*.jpg*")))
    # png and jpg image formats are accepted
    my_image = PhotoImage(file=imagefile)
    position = inputTxt.index(INSERT)
    inputTxt.image_create(position, image=my_image)

    # print(imagefile)

# function to make text italics

def font_italics():
    italic_font = font.Font(inputTxt, inputTxt.cget("font"))
    italic_font.configure(slant="italic")
    inputTxt.tag_configure("italic", font=italic_font)
    current_tags = inputTxt.tag_names("sel.first")
    if "italic" in current_tags:
        # remove italic
        inputTxt.tag_remove("italic", "sel.first", "sel.last")

    elif "bold_underline" in current_tags:
        inputTxt.tag_config("bold_italic_underline", font=(
            "Arial", 15, "bold italic underline"))
        inputTxt.tag_add("bold_italic_underline", "sel.first", "sel.last")

    elif "bold" in current_tags:
        inputTxt.tag_config("bold_italic", font=("Arial", 15, "bold italic"))
        inputTxt.tag_add("bold_italic", "sel.first", "sel.last")

    elif "underline" in current_tags:
        inputTxt.tag_config("italic_underline", font=(
            "Arial", 15, "italic underline"))
        inputTxt.tag_add("italic_underline", "sel.first", "sel.last")

    else:
        # add italic
        inputTxt.tag_add("italic", "sel.first", "sel.last")

# function to make text bold

def bolder():
    bold_font = font.Font(inputTxt, inputTxt.cget("font"))
    bold_font.configure(weight="bold")
    inputTxt.tag_configure("bold", font=bold_font)
    current_tags = inputTxt.tag_names("sel.first")
    # print(current_tags)

    if "bold" in current_tags:
        # remove bold
        inputTxt.tag_remove("bold", "sel.first", "sel.last")
        # inputTxt.tag_add("italic", "sel.first", "sel.last")
        # inputTxt.tag_config("bold_italic_underline", font=("Arial", 15, "bold italic underline"))
        # inputTxt.tag_add("bold_italic_underline", "sel.first", "sel.last")
    elif "italic_underline" in current_tags:
        # print("Yes I am here")
        inputTxt.tag_config("bold_italic_underline", font=(
            "Arial", 15, "bold italic underline"))
        inputTxt.tag_add("bold_italic_underline", "sel.first", "sel.last")

    elif "italic" in current_tags:
        inputTxt.tag_config("bold_italic", font=("Arial", 15, "bold italic"))
        inputTxt.tag_add("bold_italic", "sel.first", "sel.last")

    elif "underline" in current_tags:
        inputTxt.tag_config("bold_underline", font=(
            "Arial", 15, "bold underline"))
        inputTxt.tag_add("bold_underline", "sel.first", "sel.last")

    else:
        # add bold
        inputTxt.tag_add("bold", "sel.first", "sel.last")
        # inputTxt.tag_add("italic", "sel.first", "sel.last")

# function to add underline to font

def uline():
    ufont = font.Font(inputTxt, inputTxt.cget("font"))
    ufont.configure(underline=True)
    inputTxt.tag_configure("underline", font=ufont)
    current_tags = inputTxt.tag_names("sel.first")
    if "underline" in current_tags:
        # remove underline
        inputTxt.tag_remove("underline", "sel.first", "sel.last")

    elif "bold_italic" in current_tags:
        inputTxt.tag_config("bold_italic_underline", font=(
            "Arial", 15, "bold italic underline"))
        inputTxt.tag_add("bold_italic_underline", "sel.first", "sel.last")

    elif "bold" in current_tags:
        inputTxt.tag_config("bold_underline", font=(
            "Arial", 15, "bold underline"))
        inputTxt.tag_add("bold_underline", "sel.first", "sel.last")

    elif "italic" in current_tags:
        inputTxt.tag_config("italic_underline", font=(
            "Arial", 15, "italic underline"))
        inputTxt.tag_add("italic_underline", "sel.first", "sel.last")

    else:
        # add underline
        inputTxt.tag_add("underline", "sel.first", "sel.last")


# rightframe starts here
AddImage = Button(rightframe, text="Insert Image", command=insert_image)
AddImage.pack(padx=60, pady=60)

edit = Entry(rightframe)
edit.pack(padx=60, pady=60)
findString = Button(rightframe, text="Find", command=find)
findString.pack(padx=60, pady=60)

edit2 = Entry(rightframe)
edit2.pack(padx=60, pady=60)
replaceString = Button(rightframe, text="Replace", command=replace_string)
replaceString.pack(padx=60, pady=60)

txtHighColor = Button(rightframe, text="Text Highlight Color")
txtHighColor.pack(padx=60, pady=60)

# rightframe ends here


# leftframe starts here
makeBold = Button(leftFrame, text="BOLD", command=bolder)
makeBold.pack(padx=60, pady=60)

makeItalic = Button(leftFrame, text="ITALICS", command=font_italics)
makeItalic.pack(padx=60, pady=60)

makeUnderline = Button(leftFrame, text="Underline", command=uline)
makeUnderline.pack(padx=60, pady=60)

options = ["Calibri", "Times New Roman", "Gadugi", "ALGERIAN", "helvetica"]
clicked = StringVar()

clicked.set("helvetica")

fontStyleDrop = OptionMenu(leftFrame, clicked, *options)
fontStyleDrop.pack(padx=60, pady=60)

changeFontStyle = Button(
    leftFrame, text="Change Font Style", command=font_style)
changeFontStyle.pack(padx=60, pady=60)

# leftframe ends here

# Clears the text box when New is clicked.


def new_option():
    inputTxt.delete("1.0", "end")


# Adding functionality of opening file explorer from Open option
def browse_files():
    filename = filedialog.askopenfilename(
        initialdir="/", title="Select a File", filetypes=(("Text files", "*.txt*"), ("PDF File", "*.pdf*"), ("DOC File", "*.docx*")))
    # print("File name is: " + filename)
    # print(filename)
    split_tup = os.path.splitext(filename)
    file_extension = split_tup[1]
    # print(file_extension)
    textToInsert = ""

    # If its a .txt file then function for  opening it
    if (file_extension == ".txt"):
        with open(filename, "r") as f:
            ans = f.read()
            textToInsert += ans
        # print(textToInsert)
        # If already some text is present then remove it and open new file
        inputTxt.delete("1.0", "end")
        inputTxt.insert(END, textToInsert)

    # If its a pdf file then function for the same is as follows
    elif (file_extension == ".pdf"):
        pdfFile = open(filename, 'rb')
        readpdf = PyPDF2.PdfReader(pdfFile)
        totalpages = len(readpdf.pages)
        with pdfplumber.open(filename) as f:
            for j in range(totalpages):
                data = f.pages[j]
                textToInsert += data.extract_text()
                # print(data.extract_text())
                inputTxt.delete("1.0", "end")
                inputTxt.insert(END, textToInsert)

    elif (file_extension == ".docx"):
        doc = docx.Document(filename)
        content = doc.paragraphs
        # inputTxt.delete("1.0", "end")
        for para in content:
            # content = content + para
            # print(para.text)
            inputTxt.insert(END, para.text)
            inputTxt.insert(END, "\n")

    # If user selected a file other than .txt or .pdf
    else:
        inputTxt.delete("1.0", "end")
        inputTxt.insert(
            END, "Dear User, Please enter File with valid fle extension!!!!!\n\nPlease select .txt, .pdf, or .docx file!!!!!\n\n")

# Implementing Save As function for .txt file

def save():
    str1 = inputTxt.get("1.0", "end-1c")  # Getting the text in the text widget
    files = [('Text Document', '*.txt')]
    file = asksaveasfile(filetypes=files, defaultextension=files)
    file.write(str1)
    file.close()  # Closing the file descriptor
    # print("File is: "+str(file))

# Implementing Save As function for .docx file

def save_as_docx():
    str2 = inputTxt.get("1.0", "end-1c")
    doc = docx.Document()
    doc.add_heading('Microsoft Student Technical Club', 0)
    para = doc.add_paragraph().add_run(str2)
    para.font.size = Pt(12)
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    doc.save(file_path)
    # doc.save(filepath)


# Creating a file menu
menubar = Menu(root)
root.config(menu=menubar)
file_menu = Menu(menubar, tearoff=0)
file_menu.add_command(label="New", command=new_option)
file_menu.add_command(label="Open", command=browse_files)
file_menu.add_command(label="Save As txt", command=lambda: save())
file_menu.add_command(label="Save As docx", command=save_as_docx)
file_menu.add_separator()


# Adding a scroll bar
scroll_bar = Scrollbar(root, orient=VERTICAL)
scroll_bar.pack(side=RIGHT,fill=Y)
inputTxt.config(yscrollcommand=scroll_bar.set)
scroll_bar.config(command=inputTxt.yview)

# Adding exit to file menu
file_menu.add_command(label="Exit", command=root.destroy)
menubar.add_cascade(label="File", menu=file_menu)

inputTxt.pack(side=LEFT)
root.mainloop()
